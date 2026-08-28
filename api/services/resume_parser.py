"""
Resume Parser Service — CyberMentor
Supports clean text extraction from PDF (.pdf), Microsoft Word (.docx), and plain text formats.
"""

import io
import logging
from pypdf import PdfReader
import docx

logger = logging.getLogger("cybermentor.resume_parser")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract clean readable text from PDF bytes."""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        pages_text = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                pages_text.append(text.strip())
        return "\n\n".join(pages_text)
    except Exception as e:
        logger.error(f"Error parsing PDF resume: {e}")
        raise ValueError(f"Could not parse PDF document: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract clean readable text from DOCX bytes."""
    try:
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Also extract table text if present
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n\n".join(paragraphs)
    except Exception as e:
        logger.error(f"Error parsing DOCX resume: {e}")
        raise ValueError(f"Could not parse Word document: {str(e)}")


def parse_resume_bytes(file_bytes: bytes, filename: str) -> str:
    """Detect file type and extract clean plain text."""
    lower_name = filename.lower()
    if lower_name.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif lower_name.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif lower_name.endswith((".txt", ".md", ".rtf", ".json")):
        return file_bytes.decode("utf-8", errors="replace")
    else:
        # Try PDF then Word then UTF-8
        try:
            return extract_text_from_pdf(file_bytes)
        except Exception:
            try:
                return extract_text_from_docx(file_bytes)
            except Exception:
                return file_bytes.decode("utf-8", errors="replace")
