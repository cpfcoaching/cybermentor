"""
Resume Exporter Service — CyberMentor

Generates publication-quality Word (.docx) and PDF (.pdf) documents
from Markdown/plain text resumes.
"""

import io
import re
import logging
from typing import List, Tuple

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

logger = logging.getLogger("cybermentor.resume_exporter")


def _strip_markdown_links(text: str) -> str:
    """Convert [text](url) to 'text (url)' or just 'text' if standard."""
    return re.sub(r'\[([^\]]+)\]\(([^)]+)\)', r'\1 (\2)', text)


def _parse_inline_bold(text: str) -> List[Tuple[str, bool]]:
    """
    Parse a string containing **bold** markup into a list of (text_chunk, is_bold) tuples.
    """
    # Pattern to match **bold text**
    parts = []
    tokens = re.split(r'(\*\*[^*]+\*\*)', text)
    for token in tokens:
        if not token:
            continue
        if token.startswith("**") and token.endswith("**") and len(token) >= 4:
            parts.append((token[2:-2], True))
        else:
            parts.append((token, False))
    return parts


def generate_docx_resume(resume_markdown: str, filename: str = "Cybersecurity_Resume.docx") -> bytes:
    """
    Convert resume markdown text into a professionally styled DOCX document.
    """
    doc = docx.Document()

    # Set page margins to 0.65 in
    for section in doc.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)

    # Base Colors
    NAVY_HEX = "1B365D"
    CHARCOAL_HEX = "2B2D42"
    MUTED_HEX = "555555"

    NAVY_COLOR = RGBColor(0x1B, 0x36, 0x5D)
    CHARCOAL_COLOR = RGBColor(0x2B, 0x2D, 0x42)
    MUTED_COLOR = RGBColor(0x55, 0x55, 0x55)

    # Set Normal Style Font
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = CHARCOAL_COLOR

    lines = resume_markdown.strip().split("\n")
    i = 0
    total_lines = len(lines)

    def add_horizontal_line(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="6" w:space="2" w:color="{NAVY_HEX}"/></w:pBdr>')
        pPr.append(pBdr)

    while i < total_lines:
        raw_line = lines[i].strip()
        i += 1

        if not raw_line or raw_line in ["---", "***", "___"]:
            continue

        # 1. Main Candidate Name Heading (# Name)
        if raw_line.startswith("# ") and not raw_line.startswith("## "):
            name_text = raw_line[2:].strip().replace("**", "")
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(name_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = NAVY_COLOR
            continue

        # 2. Contact Header Info (e.g. phone | email | linkedin)
        if ("@" in raw_line or "linkedin" in raw_line.lower() or "|" in raw_line) and i <= 6:
            clean_contact = re.sub(r'[*_`]', '', raw_line)
            clean_contact = _strip_markdown_links(clean_contact)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(clean_contact)
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
            run.font.color.rgb = MUTED_COLOR
            continue

        # 3. Target Title Subheading (### TITLE or *Title*)
        if (raw_line.startswith("### ") or (raw_line.startswith("*") and raw_line.endswith("*"))) and i <= 8 and not any(k in raw_line.upper() for k in ["SUMMARY", "EXPERIENCE", "COMPETENCIES", "EDUCATION", "PUBLICATIONS"]):
            title_text = raw_line.replace("#", "").replace("*", "").strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(title_text)
            run.font.name = 'Calibri'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = NAVY_COLOR
            continue

        # 4. Major Section Heading (## SECTION or ### SECTION)
        clean_section = raw_line.replace("#", "").replace("*", "").strip()
        is_major_section = (
            raw_line.startswith("## ") or
            (raw_line.startswith("### ") and any(k in clean_section.upper() for k in [
                "SUMMARY", "COMPETENCIES", "EXPERIENCE", "PUBLICATIONS", "SKILLS", "EDUCATION", "CREDENTIALS", "LEADERSHIP"
            ]))
        )

        if is_major_section:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_with_next = True
            run = p.add_run(clean_section.upper())
            run.font.name = 'Calibri'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = NAVY_COLOR
            add_horizontal_line(p)
            continue

        # 5. Role / Company Subheading (#### COMPANY | Date)
        if raw_line.startswith("#### ") or (raw_line.startswith("**") and ("|" in raw_line or "–" in raw_line or " - " in raw_line or "20" in raw_line)):
            role_text = raw_line.replace("####", "").strip()
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.keep_with_next = True

            # Parse bold spans
            chunks = _parse_inline_bold(role_text)
            for chunk_text, is_bold in chunks:
                run = p.add_run(chunk_text)
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)
                run.font.bold = True
                run.font.color.rgb = CHARCOAL_COLOR
            continue

        # 6. Bullet Point
        if raw_line.startswith("* ") or raw_line.startswith("- ") or raw_line.startswith("• "):
            bullet_text = raw_line[2:].strip()
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.1

            chunks = _parse_inline_bold(bullet_text)
            for chunk_text, is_bold in chunks:
                run = p.add_run(_strip_markdown_links(chunk_text))
                run.font.name = 'Calibri'
                run.font.size = Pt(9.5)
                run.font.bold = is_bold
                run.font.color.rgb = CHARCOAL_COLOR
            continue

        # 7. Standard Paragraph / Executive Summary Text
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15

        chunks = _parse_inline_bold(raw_line)
        for chunk_text, is_bold in chunks:
            run = p.add_run(_strip_markdown_links(chunk_text))
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
            run.font.bold = is_bold
            run.font.color.rgb = CHARCOAL_COLOR

    target_stream = io.BytesIO()
    doc.save(target_stream)
    return target_stream.getvalue()


def generate_pdf_resume(resume_markdown: str, filename: str = "Cybersecurity_Resume.pdf") -> bytes:
    """
    Convert resume markdown text into a professionally styled PDF document using ReportLab.
    """
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, ListFlowable, ListItem
    from reportlab.lib.units import inch

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=0.55 * inch,
        rightMargin=0.55 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch
    )

    styles = getSampleStyleSheet()

    # Custom Palette
    NAVY = colors.HexColor("#1B365D")
    CHARCOAL = colors.HexColor("#2B2D42")
    MUTED = colors.HexColor("#555555")

    # Typography Styles
    title_style = ParagraphStyle(
        'ResumeTitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=17,
        leading=21,
        textColor=NAVY,
        alignment=1, # Center
        spaceAfter=2
    )

    contact_style = ParagraphStyle(
        'ResumeContact',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8.5,
        leading=11,
        textColor=MUTED,
        alignment=1, # Center
        spaceAfter=3
    )

    subtitle_style = ParagraphStyle(
        'ResumeSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=10,
        leading=13,
        textColor=NAVY,
        alignment=1, # Center
        spaceAfter=6
    )

    section_style = ParagraphStyle(
        'ResumeSection',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=14,
        textColor=NAVY,
        spaceBefore=7,
        spaceAfter=2,
        keepWithNext=True
    )

    role_style = ParagraphStyle(
        'ResumeRole',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=9.5,
        leading=12,
        textColor=CHARCOAL,
        spaceBefore=4,
        spaceAfter=1,
        keepWithNext=True
    )

    body_style = ParagraphStyle(
        'ResumeBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8.8,
        leading=11.5,
        textColor=CHARCOAL,
        spaceBefore=1,
        spaceAfter=2
    )

    bullet_style = ParagraphStyle(
        'ResumeBullet',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=8.8,
        leading=11.5,
        textColor=CHARCOAL,
        leftIndent=12,
        firstLineIndent=-8,
        spaceBefore=1,
        spaceAfter=1.5
    )

    story = []

    def md_to_reportlab_html(text: str) -> str:
        # Convert **bold** to <b>bold</b>
        text = _strip_markdown_links(text)
        text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        # Fix back bold formatting
        text = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', text)
        text = re.sub(r'\*([^*]+)\*', r'<i>\1</i>', text)
        return text

    lines = resume_markdown.strip().split("\n")
    i = 0
    total_lines = len(lines)

    while i < total_lines:
        raw_line = lines[i].strip()
        i += 1

        if not raw_line or raw_line in ["---", "***", "___"]:
            continue

        # 1. Main Candidate Name Heading (# Name)
        if raw_line.startswith("# ") and not raw_line.startswith("## "):
            name_text = raw_line[2:].strip().replace("**", "")
            story.append(Paragraph(name_text, title_style))
            continue

        # 2. Contact Header Info
        if ("@" in raw_line or "linkedin" in raw_line.lower() or "|" in raw_line) and i <= 6:
            clean_contact = re.sub(r'[*_`]', '', raw_line)
            story.append(Paragraph(md_to_reportlab_html(clean_contact), contact_style))
            continue

        # 3. Target Title Subheading
        if (raw_line.startswith("### ") or (raw_line.startswith("*") and raw_line.endswith("*"))) and i <= 8 and not any(k in raw_line.upper() for k in ["SUMMARY", "EXPERIENCE", "COMPETENCIES", "EDUCATION", "PUBLICATIONS"]):
            title_text = raw_line.replace("#", "").replace("*", "").strip()
            story.append(Paragraph(title_text, subtitle_style))
            story.append(HRFlowable(width="100%", thickness=0.75, color=NAVY, spaceBefore=2, spaceAfter=4))
            continue

        # 4. Major Section Heading
        clean_section = raw_line.replace("#", "").replace("*", "").strip()
        is_major_section = (
            raw_line.startswith("## ") or
            (raw_line.startswith("### ") and any(k in clean_section.upper() for k in [
                "SUMMARY", "COMPETENCIES", "EXPERIENCE", "PUBLICATIONS", "SKILLS", "EDUCATION", "CREDENTIALS", "LEADERSHIP"
            ]))
        )

        if is_major_section:
            story.append(Paragraph(clean_section.upper(), section_style))
            story.append(HRFlowable(width="100%", thickness=0.5, color=NAVY, spaceBefore=1, spaceAfter=3))
            continue

        # 5. Role / Company Subheading
        if raw_line.startswith("#### ") or (raw_line.startswith("**") and ("|" in raw_line or "–" in raw_line or " - " in raw_line or "20" in raw_line)):
            role_text = raw_line.replace("####", "").strip()
            story.append(Paragraph(md_to_reportlab_html(role_text), role_style))
            continue

        # 6. Bullet Point
        if raw_line.startswith("* ") or raw_line.startswith("- ") or raw_line.startswith("• "):
            bullet_text = raw_line[2:].strip()
            bullet_html = f"&bull; {md_to_reportlab_html(bullet_text)}"
            story.append(Paragraph(bullet_html, bullet_style))
            continue

        # 7. Standard Paragraph / Executive Summary
        story.append(Paragraph(md_to_reportlab_html(raw_line), body_style))

    doc.build(story)
    return buffer.getvalue()
